import streamlit as st
from docx import Document
import PyPDF2
import re
import pandas as pd
from collections import Counter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
import io

# --- CONFIG ---
st.set_page_config(page_title="ATS Expert PRO (Cloud)", layout="centered")
st.title("🧠 ATS Expert PRO (Versión Completa)")
st.markdown("---")

uploaded_files = st.file_uploader(
    "Sube múltiples CVs",
    type=["pdf", "docx"],
    accept_multiple_files=True
)

# --- DATA ---
SKILLS_CLAVE = [
    "python","sql","java","c++","javascript","typescript",
    "excel","power bi","tableau",
    "aws","azure","gcp","docker","kubernetes",
    "git","linux","api","rest",
    "machine learning","data analysis","etl","spark"
]

PALABRAS_DEBILES = ["responsable","apoyo","ayuda","encargado","participé"]

# --- FUNCIONES BASE ---
def extraer_texto(file):
    texto = ""
    if file.type == "application/pdf":
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            texto += page.extract_text() or ""
    else:
        doc = Document(file)
        for p in doc.paragraphs:
            texto += p.text + "\n"
    return texto.lower()

def limpiar_texto(texto):
    texto = re.sub(r'[^a-zA-Z0-9\s]', ' ', texto.lower())
    return re.sub(r'\s+', ' ', texto)

# --- ORTOGRAFIA BASICA ---
def analizar_ortografia(texto):
    errores = []

    if "  " in texto:
        errores.append("✍️ Doble espacio detectado")

    if re.search(r'\b(q|xk|pq)\b', texto):
        errores.append("⚠️ Uso de lenguaje informal")

    if re.search(r'\bteh\b', texto):
        errores.append("❌ Error ortográfico común detectado")

    return errores

# --- LOGROS VS RESPONSABILIDADES ---
def analizar_logros(texto):
    lineas = texto.split("\n")
    logros = 0
    responsabilidades = 0

    for l in lineas:
        if re.search(r'\d+%|\$|\d+', l):
            logros += 1
        elif any(p in l for p in PALABRAS_DEBILES):
            responsabilidades += 1

    return logros, responsabilidades

# --- BULLETS ---
def detectar_bullets(texto):
    return len(re.findall(r'[-•●]', texto))

# --- ANALISIS ATS BASE ---
def analizar(texto, vacante):
    cv_clean = limpiar_texto(texto)
    vac_clean = limpiar_texto(vacante)

    kw_cv = set(cv_clean.split())
    kw_vac = set(vac_clean.split())

    return int((len(kw_cv & kw_vac) / max(len(kw_vac),1)) * 100)

# --- SCORE AVANZADO ---
def score_avanzado(texto, vacante):
    score_kw = analizar(texto, vacante)
    errores = len(analizar_ortografia(texto))
    logros, resp = analizar_logros(texto)

    score = (
        score_kw * 0.4 +
        max(0, 100 - errores*5) * 0.2 +
        min(logros*10, 100) * 0.2 +
        (100 if logros > resp else 50) * 0.2
    )

    return int(score)

# --- AUDITORIA COMPLETA ---
def auditoria_cv(texto, vacante):
    observaciones = []

    palabras = texto.split()
    conteo = Counter(palabras)

    repetidas = [p for p, c in conteo.items() if c > 10 and len(p) > 4]
    if repetidas:
        observaciones.append(f"⚠️ Palabras repetidas: {repetidas[:5]}")

    if any(p in palabras for p in PALABRAS_DEBILES):
        observaciones.append("⚠️ Lenguaje débil detectado")

    if not re.search(r'\d+%|\$', texto):
        observaciones.append("📉 Falta de métricas cuantificables")

    if "  " in texto:
        observaciones.append("✍️ Problemas de formato")

    años = re.findall(r'(20\d{2})', texto)
    if len(años) > 1:
        años_int = sorted([int(a) for a in años])
        gaps = [años_int[i+1] - años_int[i] for i in range(len(años_int)-1)]
        if any(g > 3 for g in gaps):
            observaciones.append("📅 Posibles gaps laborales")

    kw_cv = set(limpiar_texto(texto).split())
    kw_vac = set(limpiar_texto(vacante).split())

    faltantes = kw_vac - kw_cv
    if len(faltantes) > 20:
        observaciones.append("🎯 Baja coincidencia ATS")

    skills_faltantes = [s for s in SKILLS_CLAVE if s in vacante and s not in texto]
    if skills_faltantes:
        observaciones.append(f"🧠 Skills faltantes: {skills_faltantes[:5]}")

    return observaciones

# --- REESCRITURA AUTOMATICA ---
def reescribir_cv(texto):
    mejoras = []

    for linea in texto.split("\n"):
        nueva = linea

        if "responsable de" in linea:
            nueva = linea.replace("responsable de", "Lideré")

        if "apoyo en" in linea:
            nueva = nueva.replace("apoyo en", "Contribuí a")

        if "encargado de" in linea:
            nueva = nueva.replace("encargado de", "Gestioné")

        mejoras.append(nueva)

    return "\n".join(mejoras)

# --- PDF EN MEMORIA ---
def generar_pdf_bytes(nombre, score, observaciones, mejoras):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer)
    styles = getSampleStyleSheet()

    contenido = []
    contenido.append(Paragraph(f"Reporte ATS - {nombre}", styles['Title']))
    contenido.append(Spacer(1, 10))
    contenido.append(Paragraph(f"Score: {score}", styles['Heading2']))

    contenido.append(Paragraph("Observaciones:", styles['Heading3']))
    for o in observaciones:
        contenido.append(Paragraph(o, styles['Normal']))

    contenido.append(Spacer(1, 10))
    contenido.append(Paragraph("CV Mejorado:", styles['Heading3']))

    for m in mejoras.split("\n")[:50]:
        contenido.append(Paragraph(m, styles['Normal']))

    doc.build(contenido)
    buffer.seek(0)

    return buffer

# --- DOCX EN MEMORIA ---
def generar_docx_bytes(texto_mejorado):
    buffer = io.BytesIO()
    doc = Document()
    doc.add_heading('CV Optimizado', 0)

    for linea in texto_mejorado.split("\n"):
        doc.add_paragraph(linea)

    doc.save(buffer)
    buffer.seek(0)

    return buffer

# --- UI PRINCIPAL ---
if uploaded_files:
    vacante = st.text_area("Pega la descripción del puesto")

    if st.button("🚀 Analizar candidatos"):
        resultados = []

        for file in uploaded_files:
            texto = extraer_texto(file)

            if texto.strip():
                score = score_avanzado(texto, vacante)
                obs = auditoria_cv(texto, vacante)
                errores = analizar_ortografia(texto)
                logros, resp = analizar_logros(texto)
                bullets = detectar_bullets(texto)
                texto_mejorado = reescribir_cv(texto)

                resultados.append({
                    "Archivo": file.name,
                    "Score": score,
                    "Logros": logros,
                    "Responsabilidades": resp,
                    "Bullets": bullets,
                    "Errores": len(errores)
                })

        df = pd.DataFrame(resultados).sort_values(by="Score", ascending=False)

        st.subheader("🏆 Ranking")
        st.dataframe(df)

        csv = df.to_csv(index=False).encode("utf-8")
        st.download_button("📄 Descargar CSV", csv, "reporte_cv.csv")

        st.markdown("## 🔍 Detalle por candidato")

        for file in uploaded_files:
            texto = extraer_texto(file)

            if texto.strip():
                score = score_avanzado(texto, vacante)
                obs = auditoria_cv(texto, vacante)
                errores = analizar_ortografia(texto)
                logros, resp = analizar_logros(texto)
                bullets = detectar_bullets(texto)
                texto_mejorado = reescribir_cv(texto)

                pdf_buffer = generar_pdf_bytes(file.name, score, obs + errores, texto_mejorado)
                docx_buffer = generar_docx_bytes(texto_mejorado)

                with st.expander(f"📄 {file.name} - {score}%"):

                    st.write("### 📊 Análisis")
                    st.write(f"✅ Score ATS: {score}")
                    st.write(f"🏆 Logros: {logros}")
                    st.write(f"📋 Responsabilidades: {resp}")
                    st.write(f"🔹 Bullets detectados: {bullets}")
                    st.write(f"❌ Errores: {len(errores)}")

                    st.write("### ⚠️ Observaciones")
                    for o in obs:
                        st.write(o)

                    st.write("### ✍️ Errores detectados")
                    for e in errores:
                        st.write(e)

                    st.write("### 🚀 CV Reescrito (preview)")
                    st.text(texto_mejorado[:1000])

                    st.download_button(
                        label="📄 Descargar PDF reporte",
                        data=pdf_buffer,
                        file_name=f"{file.name}_reporte.pdf",
                        mime="application/pdf"
                    )

                    st.download_button(
                        label="📝 Descargar CV optimizado (DOCX)",
                        data=docx_buffer,
                        file_name=f"{file.name}_mejorado.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

else:
    st.info("Sube CVs para comenzar")

st.markdown("---")
st.caption("ATS Expert PRO - Versión Completa Funcional")